import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Artifact Paths
ARTIFACT_DIR = r"C:\Users\harsh\.gemini\antigravity\brain\72523d82-b17d-48ea-8246-7dc9bd90f81b"
FILES_TO_EXPORT = [
    "project_roadmap.md",
    "implementation_plan.md",
    "walkthrough.md"
]
OUTPUT_FILE = "Project_Documentation.docx"

def add_markdown_content(doc, filepath):
    filename = os.path.basename(filepath)
    
    # Add Title for this section
    doc.add_page_break()
    title = doc.add_heading(filename, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    
    for line in lines:
        line = line.rstrip()
        
        # Code Blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # List Items
        elif line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
             p = doc.add_paragraph(line.strip())
             p.style = 'List Bullet'
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(line.strip()[2:])
            p.style = 'List Bullet'
        
        # Normal Text
        else:
            if line.strip():
                doc.add_paragraph(line)

def create_docx():
    doc = Document()
    
    # Main Title
    title = doc.add_heading('Law Firm Connect - Project Documentation', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f'Generated on {os.popen("date /t").read().strip()}')
    
    for filename in FILES_TO_EXPORT:
        filepath = os.path.join(ARTIFACT_DIR, filename)
        if os.path.exists(filepath):
            print(f"Adding {filename}...")
            add_markdown_content(doc, filepath)
        else:
            print(f"Warning: {filename} not found at {filepath}")
            
    doc.save(OUTPUT_FILE)
    print(f"Successfully created {OUTPUT_FILE}")

if __name__ == "__main__":
    try:
        create_docx()
    except ImportError:
        print("python-docx not found. Installing...")
        os.system("pip install python-docx")
        create_docx()
